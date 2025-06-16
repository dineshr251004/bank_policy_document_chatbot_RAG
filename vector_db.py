from langchain_ollama import OllamaEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from docx import Document as DocxDocument  
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os

DOCX_FOLDER = "./synthetic_policy_documents"
DB_LOCATION = "./chroma_bank_policy_db"

embeddings = OllamaEmbeddings(model="mxbai-embed-large")

add_documents = not os.path.exists(DB_LOCATION)

if add_documents:
    documents = []
    ids = []
    doc_id = 0
    
    for filename in os.listdir(DOCX_FOLDER):
        if filename.endswith(".docx"):
            filepath = os.path.join(DOCX_FOLDER, filename)
            docx_file = DocxDocument(filepath)
            full_text = "\n".join([para.text for para in docx_file.paragraphs if para.text.strip()])

            splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
            chunks = splitter.split_text(full_text)
            
            for i, chunk in enumerate(chunks):
                document = Document(
                    page_content=chunk,
                    metadata={"source_file": filename, "chunk_id": i},
                    id=str(doc_id)
                )
                documents.append(document)
                ids.append(str(doc_id))
                doc_id += 1

vector_store = Chroma(
    collection_name="bank_policy_docs",
    persist_directory=DB_LOCATION,
    embedding_function=embeddings
)

if add_documents:
    vector_store.add_documents(documents=documents, ids=ids)

retriever = vector_store.as_retriever(search_kwargs={"k": 5})
